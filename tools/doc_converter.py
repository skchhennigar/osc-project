#!/usr/bin/env python3
"""
OSC-Proj Document Converter
Converts various document formats to markdown for Claude Code processing
"""

import os
import sys
import argparse
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import json

# Import libraries for different document types
try:
    import docx
    from docx.document import Document as DocxDocument
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    import mammoth
except ImportError:
    mammoth = None

import re
from html import unescape
import markdown

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentConverter:
    """Convert various document formats to markdown"""
    
    def __init__(self, output_dir: str = "converted_docs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Supported file extensions
        self.supported_formats = {
            '.docx': self._convert_docx,
            '.doc': self._convert_docx,
            '.xlsx': self._convert_excel,
            '.xls': self._convert_excel,
            '.pdf': self._convert_pdf,
            '.txt': self._convert_text,
            '.md': self._copy_markdown,
            '.html': self._convert_html,
            '.rtf': self._convert_rtf,
            '.csv': self._convert_csv
        }
    
    def convert_file(self, file_path: str, output_name: Optional[str] = None) -> Dict[str, Any]:
        """Convert a single file to markdown"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Generate output filename
        if output_name:
            output_file = self.output_dir / f"{output_name}.md"
        else:
            output_file = self.output_dir / f"{file_path.stem}.md"
        
        logger.info(f"Converting {file_path} to {output_file}")
        
        # Convert using appropriate method
        converter = self.supported_formats[extension]
        result = converter(file_path, output_file)
        
        # Add metadata
        result.update({
            'input_file': str(file_path),
            'output_file': str(output_file),
            'conversion_time': datetime.now().isoformat(),
            'file_size': file_path.stat().st_size
        })
        
        return result
    
    def convert_directory(self, directory: str, recursive: bool = True) -> List[Dict[str, Any]]:
        """Convert all supported files in a directory"""
        directory = Path(directory)
        results = []
        
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    result = self.convert_file(file_path)
                    results.append(result)
                except Exception as e:
                    logger.error(f"Failed to convert {file_path}: {e}")
                    results.append({
                        'input_file': str(file_path),
                        'error': str(e),
                        'status': 'failed'
                    })
        
        return results
    
    def _convert_docx(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert Word document to markdown"""
        if docx is None:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            # Try mammoth for better HTML conversion if available
            if mammoth:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html_content = result.html
                    markdown_content = self._html_to_markdown(html_content)
            else:
                # Fallback to basic docx extraction
                doc = docx.Document(file_path)
                markdown_content = self._docx_to_markdown(doc)
            
            # Add document metadata
            metadata = self._extract_docx_metadata(file_path)
            full_content = self._add_frontmatter(markdown_content, metadata)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                'status': 'success',
                'content_type': 'document',
                'metadata': metadata,
                'word_count': len(markdown_content.split())
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert DOCX: {e}")
    
    def _convert_excel(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert Excel file to markdown"""
        if pd is None:
            raise ImportError("pandas not installed. Run: pip install pandas openpyxl")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            markdown_content = f"# {file_path.stem}\n\n"
            
            sheet_info = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Skip empty sheets
                if df.empty:
                    continue
                
                markdown_content += f"## {sheet_name}\n\n"
                
                # Convert to markdown table
                markdown_table = df.to_markdown(index=False)
                markdown_content += markdown_table + "\n\n"
                
                sheet_info.append({
                    'name': sheet_name,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'columns_list': df.columns.tolist()
                })
            
            metadata = {
                'file_type': 'spreadsheet',
                'sheets': sheet_info,
                'total_sheets': len(sheet_info)
            }
            
            full_content = self._add_frontmatter(markdown_content, metadata)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                'status': 'success',
                'content_type': 'spreadsheet',
                'metadata': metadata,
                'sheets_processed': len(sheet_info)
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert Excel: {e}")
    
    def _convert_pdf(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert PDF to markdown"""
        markdown_content = ""
        page_count = 0
        
        try:
            # Try pdfplumber first (better text extraction)
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text:
                            markdown_content += f"## Page {i}\n\n{text}\n\n"
            
            # Fallback to PyPDF2
            elif PyPDF2:
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    page_count = len(pdf_reader.pages)
                    
                    for i, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            markdown_content += f"## Page {i}\n\n{text}\n\n"
            else:
                raise ImportError("No PDF library available. Install: pip install pdfplumber")
            
            if not markdown_content.strip():
                markdown_content = f"# {file_path.stem}\n\n*PDF content could not be extracted as text*\n"
            
            metadata = {
                'file_type': 'pdf',
                'pages': page_count,
                'extracted_text': bool(markdown_content.strip())
            }
            
            full_content = self._add_frontmatter(markdown_content, metadata)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                'status': 'success',
                'content_type': 'pdf',
                'metadata': metadata,
                'pages_processed': page_count
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert PDF: {e}")
    
    def _convert_text(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert plain text to markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Basic text-to-markdown conversion
            markdown_content = f"# {file_path.stem}\n\n{content}"
            
            metadata = {
                'file_type': 'text',
                'encoding': 'utf-8'
            }
            
            full_content = self._add_frontmatter(markdown_content, metadata)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                'status': 'success',
                'content_type': 'text',
                'metadata': metadata,
                'character_count': len(content)
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert text: {e}")
    
    def _convert_csv(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert CSV to markdown table"""
        if pd is None:
            raise ImportError("pandas not installed. Run: pip install pandas")
        
        try:
            df = pd.read_csv(file_path)
            
            markdown_content = f"# {file_path.stem}\n\n"
            markdown_table = df.to_markdown(index=False)
            markdown_content += markdown_table
            
            metadata = {
                'file_type': 'csv',
                'rows': len(df),
                'columns': len(df.columns),
                'columns_list': df.columns.tolist()
            }
            
            full_content = self._add_frontmatter(markdown_content, metadata)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            return {
                'status': 'success',
                'content_type': 'csv',
                'metadata': metadata,
                'rows_processed': len(df)
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert CSV: {e}")
    
    def _copy_markdown(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Copy existing markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'status': 'success',
                'content_type': 'markdown',
                'action': 'copied'
            }
            
        except Exception as e:
            raise Exception(f"Failed to copy markdown: {e}")
    
    def _convert_html(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Convert HTML to markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            markdown_content = self._html_to_markdown(html_content)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return {
                'status': 'success',
                'content_type': 'html',
                'conversion': 'html_to_markdown'
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert HTML: {e}")
    
    def _convert_rtf(self, file_path: Path, output_file: Path) -> Dict[str, Any]:
        """Basic RTF to text conversion"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            
            # Very basic RTF stripping - remove RTF codes
            text_content = re.sub(r'\{[^}]*\}', '', rtf_content)
            text_content = re.sub(r'\\[a-z]+\d*', '', text_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            markdown_content = f"# {file_path.stem}\n\n{text_content}"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return {
                'status': 'success',
                'content_type': 'rtf',
                'note': 'Basic conversion - formatting may be lost'
            }
            
        except Exception as e:
            raise Exception(f"Failed to convert RTF: {e}")
    
    def _docx_to_markdown(self, doc: DocxDocument) -> str:
        """Convert DOCX document to markdown using python-docx"""
        markdown_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_lines.append("")
                continue
            
            # Handle different paragraph styles
            style = paragraph.style.name.lower()
            
            if 'heading 1' in style:
                markdown_lines.append(f"# {text}")
            elif 'heading 2' in style:
                markdown_lines.append(f"## {text}")
            elif 'heading 3' in style:
                markdown_lines.append(f"### {text}")
            elif 'heading 4' in style:
                markdown_lines.append(f"#### {text}")
            elif 'heading 5' in style:
                markdown_lines.append(f"##### {text}")
            elif 'heading 6' in style:
                markdown_lines.append(f"###### {text}")
            else:
                markdown_lines.append(text)
        
        # Handle tables
        for table in doc.tables:
            markdown_lines.append("")  # Add space before table
            
            # Header row
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
            
            # Data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |")
            
            markdown_lines.append("")  # Add space after table
        
        return "\n".join(markdown_lines)
    
    def _html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to markdown"""
        # Basic HTML to markdown conversion
        # Remove HTML tags and convert common elements
        
        # Replace common HTML elements
        html_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', html_content)
        html_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', html_content)
        html_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', html_content)
        html_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1', html_content)
        html_content = re.sub(r'<h5[^>]*>(.*?)</h5>', r'##### \1', html_content)
        html_content = re.sub(r'<h6[^>]*>(.*?)</h6>', r'###### \1', html_content)
        
        html_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', html_content)
        html_content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', html_content)
        html_content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', html_content)
        html_content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', html_content)
        
        html_content = re.sub(r'<p[^>]*>', '\n', html_content)
        html_content = re.sub(r'</p>', '\n', html_content)
        html_content = re.sub(r'<br[^>]*>', '\n', html_content)
        
        # Remove all other HTML tags
        html_content = re.sub(r'<[^>]+>', '', html_content)
        
        # Unescape HTML entities
        html_content = unescape(html_content)
        
        # Clean up whitespace
        html_content = re.sub(r'\n\s*\n', '\n\n', html_content)
        
        return html_content.strip()
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        if docx is None:
            return {'file_type': 'document'}
        
        try:
            doc = docx.Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                'file_type': 'document',
                'title': core_props.title or file_path.stem,
                'author': core_props.author,
                'subject': core_props.subject,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by
            }
            
            # Remove None values
            return {k: v for k, v in metadata.items() if v is not None}
            
        except Exception:
            return {'file_type': 'document', 'title': file_path.stem}
    
    def _add_frontmatter(self, content: str, metadata: Dict[str, Any]) -> str:
        """Add YAML frontmatter to markdown content"""
        frontmatter_lines = ["---"]
        
        for key, value in metadata.items():
            if isinstance(value, list):
                frontmatter_lines.append(f"{key}: {json.dumps(value)}")
            elif isinstance(value, str):
                frontmatter_lines.append(f"{key}: {json.dumps(value)}")
            else:
                frontmatter_lines.append(f"{key}: {value}")
        
        frontmatter_lines.extend([
            f"converted_date: {datetime.now().isoformat()}",
            "status: converted",
            "needs_review: true",
            "---",
            ""
        ])
        
        return "\n".join(frontmatter_lines) + content


def main():
    """Command line interface"""
    parser = argparse.ArgumentParser(description="Convert documents to markdown for OSC-Proj")
    parser.add_argument("input", help="Input file or directory")
    parser.add_argument("-o", "--output", default="converted_docs", help="Output directory")
    parser.add_argument("-r", "--recursive", action="store_true", help="Process directories recursively")
    parser.add_argument("-n", "--name", help="Output filename (for single file conversion)")
    parser.add_argument("-v", "--verbose", action="store_true", help="Verbose logging")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    converter = DocumentConverter(args.output)
    
    input_path = Path(args.input)
    
    try:
        if input_path.is_file():
            result = converter.convert_file(args.input, args.name)
            print(f"✅ Converted: {result['input_file']} -> {result['output_file']}")
            if result.get('error'):
                print(f"❌ Error: {result['error']}")
        
        elif input_path.is_dir():
            results = converter.convert_directory(args.input, args.recursive)
            
            successful = [r for r in results if r.get('status') == 'success']
            failed = [r for r in results if r.get('status') == 'failed']
            
            print(f"✅ Successfully converted: {len(successful)} files")
            print(f"❌ Failed to convert: {len(failed)} files")
            
            if failed:
                print("\nFailed files:")
                for result in failed:
                    print(f"  - {result['input_file']}: {result['error']}")
        
        else:
            print(f"❌ Input path does not exist: {args.input}")
            sys.exit(1)
    
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()